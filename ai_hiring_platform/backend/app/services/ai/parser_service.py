import fitz  # PyMuPDF
import docx
from app.core.logging import logger
from app.core.exceptions import UploadError

def parse_pdf(file_path: str) -> str:
    """
    Extracts raw text from a PDF document, injecting page boundaries for downstream tracking.
    """
    try:
        logger.info(f"Opening PDF document: {file_path}")
        doc = fitz.open(file_path)
        pages_content = []
        for i, page in enumerate(doc):
            page_text = page.get_text()
            # Inject page boundaries so chunking/structuring services can trace page numbers
            pages_content.append(f"--- PAGE {i + 1} ---\n{page_text}")
        
        full_text = "\n".join(pages_content)
        logger.info(f"Successfully parsed PDF: {len(full_text)} characters extracted across {len(doc)} pages.")
        return full_text
    except Exception as e:
        logger.error(f"Error parsing PDF document {file_path}: {e}", exc_info=True)
        raise UploadError(f"Failed to parse PDF document: {str(e)}")

def parse_docx(file_path: str) -> str:
    """
    Extracts raw text from a DOCX document.
    """
    try:
        logger.info(f"Opening DOCX document: {file_path}")
        doc = docx.Document(file_path)
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
        
        # Also grab text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
                    
        full_text = "\n".join(paragraphs)
        logger.info(f"Successfully parsed DOCX: {len(full_text)} characters extracted.")
        return full_text
    except Exception as e:
        logger.error(f"Error parsing DOCX document {file_path}: {e}", exc_info=True)
        raise UploadError(f"Failed to parse DOCX document: {str(e)}")
