import os
import pytest
import fitz
import docx
from app.services.ai.parser_service import parse_pdf, parse_docx
from app.services.ai.resume_structuring_service import structure_resume_to_nodes
from app.services.ai.jd_requirement_extractor import extract_requirements
from app.services.ai.vector_store_service import save_nodes_to_index, load_index
from app.services.ai.retrieval_service import retrieve_evidence_for_requirements

@pytest.fixture(scope="module")
def generate_sample_files():
    # 1. Create a sample PDF resume using fitz
    pdf_path = "./test_temp_resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    # Write some resume-style contents
    page.insert_text((50, 50), "John Doe Resume\nSkills\nPython, TensorFlow, Docker, AWS\nExperience\nDeveloped NLP systems using PyTorch.")
    doc.save(pdf_path)
    doc.close()
    
    # 2. Create a sample DOCX job description using docx
    docx_path = "./test_temp_jd.docx"
    doc_word = docx.Document()
    doc_word.add_paragraph("We are looking for a Software Engineer.")
    doc_word.add_paragraph("Requirements:")
    doc_word.add_paragraph("- Deep knowledge of Python and TensorFlow")
    doc_word.add_paragraph("- Kubernetes experience is nice to have")
    doc_word.save(docx_path)
    
    yield pdf_path, docx_path
    
    # Cleanup files
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    if os.path.exists(docx_path):
        os.remove(docx_path)

def test_document_parsers(generate_sample_files):
    pdf_path, docx_path = generate_sample_files
    
    # Verify PDF parsing
    pdf_text = parse_pdf(pdf_path)
    assert "John Doe" in pdf_text
    assert "TensorFlow" in pdf_text
    assert "PAGE 1" in pdf_text
    
    # Verify DOCX parsing
    docx_text = parse_docx(docx_path)
    assert "Software Engineer" in docx_text
    assert "TensorFlow" in docx_text

def test_resume_structuring(generate_sample_files):
    pdf_path, _ = generate_sample_files
    pdf_text = parse_pdf(pdf_path)
    
    nodes = structure_resume_to_nodes(
        text=pdf_text,
        candidate_id=99,
        resume_id=99,
        filename="test_temp_resume.pdf"
    )
    
    assert len(nodes) > 0
    # Let's inspect node metadata
    first_node = nodes[0]
    assert first_node.metadata["candidate_id"] == 99
    assert first_node.metadata["resume_id"] == 99
    assert first_node.metadata["filename"] == "test_temp_resume.pdf"
    assert "section" in first_node.metadata

def test_jd_requirement_extractor(generate_sample_files):
    _, docx_path = generate_sample_files
    jd_text = parse_docx(docx_path)
    
    requirements = extract_requirements(jd_text)

    # python and tensorflow are in taxonomy and present in text
    assert "python" in requirements
    assert "tensorflow" in requirements


def test_jd_extractor_generalizes_beyond_taxonomy():
    """Requirements outside the known taxonomy (tools, domain, soft skills, seniority)
    must now be extracted deterministically — the core RAG generalization fix."""
    jd = (
        "Requirements:\n"
        "- 5+ years of experience with data pipelines\n"
        "- Strong proficiency in Snowflake and Apache Airflow\n"
        "- Experience with stakeholder management\n"
        "- Familiarity with Terraform\n"
    )
    reqs = [r.lower() for r in extract_requirements(jd)]

    assert "terraform" in reqs                     # known taxonomy term still found
    assert "snowflake" in reqs                      # unknown tool, previously invisible
    assert "apache airflow" in reqs                 # unknown multi-word tool
    assert "stakeholder management" in reqs         # soft/domain requirement
    assert any("year" in r for r in reqs)           # seniority requirement captured
    # cue words themselves must never become requirements
    assert "familiarity" not in reqs
    assert "experience" not in reqs


def test_technical_specificity_is_generic():
    """The tech-specificity ranking signal generalizes (no hardcoded keyword list):
    a tech-dense chunk scores higher than prose, including for unseen technologies."""
    from app.services.ai.retrieval_service import _technical_specificity

    techy = _technical_specificity("Built ETL on Snowflake and Airflow with dbt models and S3.")
    prose = _technical_specificity("Collaborated closely with the team to deliver on time.")
    assert techy > prose
    assert 0 <= prose <= 100 and 0 <= techy <= 100

def test_vector_store_and_retrieval(generate_sample_files):
    pdf_path, _ = generate_sample_files
    pdf_text = parse_pdf(pdf_path)
    
    nodes = structure_resume_to_nodes(
        text=pdf_text,
        candidate_id=99,
        resume_id=99,
        filename="test_temp_resume.pdf"
    )
    
    # We must mock or execute embedding generation
    # BAAI/bge-small-en-v1.5 embedding step
    from app.services.ai.embedding_service import generate_embeddings_for_nodes
    nodes_embedded = generate_embeddings_for_nodes(nodes)
    
    # Save to FAISS vector index
    save_nodes_to_index(nodes_embedded, resume_id=99)
    
    # Load the index
    index = load_index(resume_id=99)
    assert index is not None
    
    # Retrieve evidence for requirements
    retrieval_results = retrieve_evidence_for_requirements(
        index=index,
        requirements=["python", "tensorflow"],
        top_k=2
    )
    
    assert len(retrieval_results) == 2
    # Verify structure of matches
    python_match = [r for r in retrieval_results if r["requirement"] == "python"][0]
    assert len(python_match["matches"]) > 0
    assert "score" in python_match["matches"][0]
    assert "chunk" in python_match["matches"][0]
    
    # Clean up vector index storage files created during test
    persist_dir = f"./storage/vectors/resume_99"
    if os.path.exists(persist_dir):
        for f in os.listdir(persist_dir):
            os.remove(os.path.join(persist_dir, f))
        os.rmdir(persist_dir)

def test_min_similarity_filters_weak_matches(generate_sample_files):
    """
    A requirement with no genuinely-relevant chunk must yield zero matches once the
    similarity floor is applied, and downstream must report it as "Missing" rather
    than a weak "Partial" backed by an unrelated nearest-neighbour chunk.
    """
    pdf_path, _ = generate_sample_files
    pdf_text = parse_pdf(pdf_path)

    nodes = structure_resume_to_nodes(
        text=pdf_text,
        candidate_id=98,
        resume_id=98,
        filename="test_temp_resume.pdf"
    )

    from app.services.ai.embedding_service import generate_embeddings_for_nodes
    from app.services.ai.evaluation_service import run_mock_evaluation

    nodes_embedded = generate_embeddings_for_nodes(nodes)
    save_nodes_to_index(nodes_embedded, resume_id=98)
    index = load_index(resume_id=98)
    assert index is not None

    # "kubernetes" is absent from the resume. FAISS still returns nearest neighbours
    # (e.g. the Docker/AWS chunk), so a floor is required to reject them.
    unrelated = ["kubernetes"]

    # No floor -> nearest neighbours leak through as (misleading) evidence.
    unfiltered = retrieve_evidence_for_requirements(
        index=index, requirements=unrelated, top_k=3, min_similarity=0.0
    )
    assert len(unfiltered[0]["matches"]) > 0

    # High floor -> weak matches are dropped, leaving no evidence.
    filtered = retrieve_evidence_for_requirements(
        index=index, requirements=unrelated, top_k=3, min_similarity=0.99
    )
    assert filtered[0]["matches"] == []

    # Priority 3: empty matches must surface as a genuine "Missing", not "Partial".
    report = run_mock_evaluation({"analysis_id": 1, "retrieval_results": filtered})
    kubernetes_fit = [r for r in report.requirements if r.requirement == "kubernetes"][0]
    assert kubernetes_fit.status == "Missing"
    assert "kubernetes" in report.missing_skills

    # Clean up vector index storage files created during test
    persist_dir = f"./storage/vectors/resume_98"
    if os.path.exists(persist_dir):
        for f in os.listdir(persist_dir):
            os.remove(os.path.join(persist_dir, f))
        os.rmdir(persist_dir)
